from ClassProject.CloseWs import CloseWs
from ClassProject.DataExcel import DataExcel
from ClassProject.DataExcelSiscovid import DataExcelSiscovid
from ClassProject.SiscovidBrowser import SiscovidBrowser
from ClassProject.DataToRun import DataApp
from openpyxl import Workbook
import pandas as pd
import os
import glob
import matplotlib.pyplot as plt
import seaborn as sns
import logging
from PIL import Image
import locale
import docx




class RunSiscovidFromData(SiscovidBrowser,DataApp):
    
    def __init__(self,user:str,password:str,direccion_ipress:str="Av. Angamos Oeste 300, Miraflores,Lima, Lima"):
        super().__init__(user=user,password=password,direccion_ipress=direccion_ipress)
        DataApp.__init__(self)
        self.dataframe=self.frame
        self.path_report_temp=self.path_report_temp()
        self.figure_path=self.find_path_figure()
        self.template_word=self.find_path_template_word()
        self.save_word=self.create_word_path()
        self.save_excel=self.create_path_report()
        

    def create_path_report(self):
        set_day= (pd.Timestamp.now()).strftime("%d-%m-%Y-%H-%M")
        path_report=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "Results",f"REPORTE_{set_day}.xlsx")
        return path_report
    
    def create_word_path(self):
        set_day= (pd.Timestamp.now()).strftime("%d-%m-%Y-%H-%M")
        path_word=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "Results",f"REPORTE_{set_day}.docx")
        return path_word
    
    def path_report_temp(self):
        path_report=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "Temp","REPORTE_TEMPORAL.xlsx")
        return path_report
    
    def find_path_template_word(self):
        path=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "template","template_reporte.docx")
        return path

    def find_path_figure(self):
        path=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "template","figure.png")
        return path
    
    def Ingresar_paciente(self,estado_siscovid:str,tipo_documento:str, num_documento:str,
                          Fecha,Hora_Ejecucion_de_la_prueba,Tiene_Sintomas:str,Fecha_de_inicio_de_Sintomas:str,
                                 Marque_los_Sintomas_presenta:str,Otros_especificar:str,Clasifica_clinica_severidad:str,
                                 Condicion_de_la_Persona:str,Tipo_Muestra:str,Resultado_de_la_prueba:str,OBSERVACION:str,
                            Nombre:str, Apellido_Paterno:str, Apellido_Materno:str,
                            Fecha_de_Nacimiento:str, Sexo:str, Etnia:str, Tipo_Seguro:str,
                            Procedencia_pais:str, Codigo_Pais:str, Celular:str,Correo:str,Tipo_de_Residencia:str, Direccion:str,
                            Departamento:str, Provincia:str, Distrito:str
                          ):
        
        if ((estado_siscovid=="NO INGRESADO") | (estado_siscovid=="")):
            try:
                self.buscar_paciente(tipo_documento=tipo_documento,num_documento=num_documento)
                
                if self.fix_find_pacient():
                    
                    if self.is_complete_data_paciente():
                        self.llenar_formulario_prueba(Fecha,Hora_Ejecucion_de_la_prueba,Tiene_Sintomas,Fecha_de_inicio_de_Sintomas,
                                        Marque_los_Sintomas_presenta,Otros_especificar,Clasifica_clinica_severidad,
                                        Condicion_de_la_Persona,Tipo_Muestra,Resultado_de_la_prueba,OBSERVACION)
                        estado_siscovid="ENVIADO"
                        return {"estado_siscovid":estado_siscovid, "complete":True}
                    else:
                        self.llenar_datos_paciente(Nombre, Apellido_Paterno, Apellido_Materno,
                                                    Fecha_de_Nacimiento, Sexo, Etnia, Tipo_Seguro,
                                                    Procedencia_pais, Codigo_Pais, Celular,Correo,Tipo_de_Residencia, Direccion,
                                                    Departamento, Provincia, Distrito)
                        
                        if self.is_complete_data_paciente():
                            self.llenar_formulario_prueba(Fecha,Hora_Ejecucion_de_la_prueba,Tiene_Sintomas,Fecha_de_inicio_de_Sintomas,
                                        Marque_los_Sintomas_presenta,Otros_especificar,Clasifica_clinica_severidad,
                                        Condicion_de_la_Persona,Tipo_Muestra,Resultado_de_la_prueba,OBSERVACION)
                            estado_siscovid="ENVIADO"
                            return {"estado_siscovid":estado_siscovid, "complete":True}
                        else:
                            estado_siscovid="NO INGRESADO"
                            return {"estado_siscovid":estado_siscovid, "complete":False}
                
                else:
                    estado_siscovid="NO INGRESADO"
                    return {"estado_siscovid":estado_siscovid, "complete":False}
            
            except Exception as e:
                estado_siscovid="NO INGRESADO"
                return {"estado_siscovid":estado_siscovid, "complete":False}
        
        else:
            return {"estado_siscovid":estado_siscovid, "complete":True}
    
    def Ingresar_data_to_siscovid(self):
        try:
            self.Login_Siscovid()
        except Exception as e:
            try:
                self.Login_Siscovid()
            except Exception as e:
                raise e
        
        logging.basicConfig(level=logging.INFO)   
        
        for index,row in self.dataframe.iterrows():
            estado_siscovid=row["ESTADO_SISCOVID"]
            tipo_documento=row["Tipo_Documento"]
            num_documento=row["Nro_Documento"]
            Fecha=row['Fecha']
            Hora_Ejecucion_de_la_prueba=row['Hora_Ejecucion_de_la_prueba']
            Tiene_Sintomas=row['Tiene_Sintomas']
            Fecha_de_inicio_de_Sintomas=row['Fecha_de_inicio_de_Sintomas']
            Marque_los_Sintomas_presenta=row["Marque_los_Sintomas_presenta"]
            Otros_especificar=row["Otros_especificar"]
            Clasifica_clinica_severidad=row['Clasifica_clinica_severidad']
            Condicion_de_la_Persona=row["Condicion_de_la_Persona"]
            Tipo_Muestra=row["Tipo_Muestra"]
            Resultado_de_la_prueba=row["Resultado_de_la_prueba"]
            OBSERVACION=row["OBSERVACIÓN"]
            Nombre=row["Nombre"]
            Apellido_Paterno=row["Apellido_Paterno"] 
            Apellido_Materno=row["Apellido_Materno"]
            Fecha_de_Nacimiento=row["Fecha_de_Nacimiento"]
            Sexo=row['Sexo']
            Etnia=row["Etnia"]
            Tipo_Seguro=row["Tipo_Seguro"]
            Procedencia_pais=row["Procedencia_pais"]
            Codigo_Pais=row["Codigo_Pais"]
            Celular=row["Celular"]
            Correo=row["Correo"]
            Tipo_de_Residencia=row["Tipo_de_Residencia"]
            Direccion=row["Direccion"]
            Departamento=row["Departamento"]
            Provincia=row["Provincia"]
            Distrito=row["Distrito"]
            
            dict_resultado=self.Ingresar_paciente(estado_siscovid=estado_siscovid,tipo_documento=tipo_documento
                , num_documento=num_documento,Fecha=Fecha,Hora_Ejecucion_de_la_prueba=Hora_Ejecucion_de_la_prueba,Tiene_Sintomas=Tiene_Sintomas,
                Fecha_de_inicio_de_Sintomas=Fecha_de_inicio_de_Sintomas,Marque_los_Sintomas_presenta=Marque_los_Sintomas_presenta,
                Otros_especificar=Otros_especificar,Clasifica_clinica_severidad=Clasifica_clinica_severidad,Condicion_de_la_Persona=Condicion_de_la_Persona,
                Tipo_Muestra=Tipo_Muestra,Resultado_de_la_prueba=Resultado_de_la_prueba,OBSERVACION=OBSERVACION,Nombre=Nombre,Apellido_Paterno=Apellido_Paterno,
                Apellido_Materno=Apellido_Materno,Fecha_de_Nacimiento=Fecha_de_Nacimiento,Sexo=Sexo,Etnia=Etnia,Tipo_Seguro=Tipo_Seguro,
                Procedencia_pais=Procedencia_pais,Codigo_Pais=Codigo_Pais,Celular=Celular,Correo=Correo,Tipo_de_Residencia=Tipo_de_Residencia,
                Direccion=Direccion,Departamento=Departamento,Provincia=Provincia,Distrito=Distrito)
            
            self.dataframe.at[index,"ESTADO_SISCOVID"]=dict_resultado.get("estado_siscovid")
            
            if dict_resultado.get("complete")==False:
                self.dataframe.at[index,"OBSERVACIÓN"]=self.dataframe.at[index,"OBSERVACIÓN"]+"--SCRIPT NO PUDO INGRESAR EL REGISTRO."
            if int(index)%50==0:
                CloseWs.close_excel()
                self.dataframe.to_excel(self.path_report_temp)
            
            
            logging.info(f"{index} : {tipo_documento} {num_documento} => {dict_resultado.get('estado_siscovid')} ")
        
        CloseWs.close_excel()
        self.dataframe.to_excel(self.save_word.replace(".docx",".xlsx"))
        self.dataframe.to_excel(self.path_report_temp.replace("_TEMPORAL",""))
        revisar=self.dataframe[self.dataframe['ESTADO_SISCOVID'] != 'ENVIADO'].copy()
        revisar.to_excel(self.path_report_temp.replace("REPORTE_TEMPORAL","REVISAR"))
        revisar.to_excel(self.save_excel.replace("REPORTE","REVISAR"))
        CloseWs.close_word()
        self.create_report_word()
        self.delete_old_files()
        self.reset_data()
        return True
        

    def create_report_word(self):


        # CREACIÓN GRÁFICO

        estado_siscovid = self.dataframe.groupby("ESTADO_SISCOVID")["ESTADO_SISCOVID"].count()

        # Establecer el estilo y la paleta de colores
        sns.set_style("whitegrid")
        microsoft_colors = ['#8f8f8f', '#00A4EF', '#7FBA00']
        sns.set_palette(microsoft_colors)

        # Configurar el gráfico
        fig = plt.figure(figsize=(15, 10))
        wp = {'linewidth': 1, 'edgecolor': 'white'}
        tp = {'size': 35, 'family': 'sans-serif', 'ha': 'center', 'fontstyle': 'italic', 'va': 'center', 'color': 'black'}
        index_labels = list(estado_siscovid.index)
        values = list(estado_siscovid.values)

        if len(estado_siscovid) == 1:
            plt.pie(values, labels=index_labels, autopct="%1.2f%%", startangle=20, wedgeprops=wp, textprops=tp)
        elif len(estado_siscovid) == 2:
            plt.pie(values, labels=index_labels, explode=(0, 0.2), autopct="%1.2f%%", startangle=20, wedgeprops=wp, textprops=tp)
        elif len(estado_siscovid) == 3:
            plt.pie(values, labels=index_labels, explode=(0, 0.2, 0), autopct="%1.2f%%", startangle=20, wedgeprops=wp, textprops=tp)

        plt.title('Estado de ingreso en SISCOVID', fontsize=50, family='sans-serif', ha='center', va='center', fontweight='bold')
        plt.legend(loc='lower right', labels=index_labels, prop={'size': 25, 'family': 'sans-serif'})

        fig.savefig(self.figure_path)

        def resize_png(path):
            im=Image.open(path)
            im=im.resize((500,500))
            im.save(path)

        resize_png(self.figure_path)

        #CREACIÓN WORD
        TOTAL=len(self.dataframe)

        self.dataframe["fecha"]=pd.to_datetime(self.dataframe["Fecha"], format="%d/%m/%Y")

        fecha_min=self.dataframe["fecha"].min().strftime("%d/%m/%Y")
        fecha_max=self.dataframe["fecha"].max().strftime("%d/%m/%Y")


        fecha_hoy= (pd.Timestamp.now()).strftime("%d-%m-%Y, %H:%M")



        # Establece la configuración regional en español
        locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
        date_file=(pd.Timestamp.now()).strftime("%d de %B del %Y")

        if "ENVIADO" in estado_siscovid.index:
            ENVIADOS = str(estado_siscovid.loc["ENVIADO"])
            FALTA = str(estado_siscovid.drop("ENVIADO").sum())
        else:
            ENVIADOS = str(0)
            FALTA = str(estado_siscovid.sum())
        
        if int(FALTA) > 0:
            text_1 = f"2. Se han encontrado {FALTA} registros que no han sido ingresados en SISCOVID"
            text_2 = "******* SE ADJUNTA LOS PENDIENTES EN REVISAR.xlsx ******"
        else:
            text_1="Nota:  No hay pendientes de registros para reporte manual en SISCOVID."
            text_2=""
        
        valores = { 'FECHA_MIN': fecha_min,
              'FECHA_MAX': fecha_max,
              'FECHA': fecha_hoy,
              'DATE': date_file,
              'TOTAL': str(TOTAL),
              "ENVIADOS":ENVIADOS,
              "FALTA":FALTA,
              "TEXT-1":text_1,
              "TEXT-2":text_2

              }
              
        doc = docx.Document(self.template_word)
        
       
        for parrafo in doc.paragraphs:
        # Recorre todas las runas del párrafo
            for runa in parrafo.runs:
                # Si la runa es una llave, reemplaza su valor con el objeto
                for llave, valor in valores.items():
                    runa.text = runa.text.replace("{{" + llave + "}}", valor)
            
        for paragraph in doc.paragraphs:
            if "{{IMAGEN}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{IMAGEN}}", "")
                run = paragraph.add_run()
                run.add_picture(self.figure_path, width=docx.shared.Inches(4.22), height=docx.shared.Inches(2.7))

        #GUARDAR EN RESULTS Y EN TEMP
        doc.save(self.save_word)
        doc.save(self.path_report_temp.replace("_TEMPORAL.xlsx",".docx"))
        return True


    def delete_old_files(self):
        results_path=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                              "Results")
        max_files=12
        file_list= glob.glob(os.path.join(results_path,"*"))
        if len(file_list) > max_files:
            file_list = sorted(file_list, key=os.path.getctime)
            for i in range(len(file_list) - max_files):
                os.remove(file_list[i])
    
    def reset_data(self):
        # Obtener nombre del índice y encabezados de columna
        indice_nombre = self.dataframe.index.name
        columnas_nombres = list(self.dataframe.columns)

        # Crear nuevo libro de Excel y hoja de trabajo
        wb = Workbook()
        ws = wb.active

        # Agregar encabezados a la hoja de trabajo
        ws.append([indice_nombre] + columnas_nombres)

        # Guardar archivo Excel
        wb.save(self.path)



                